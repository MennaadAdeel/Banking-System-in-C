from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

BillInfo = open("Database/Bill/Bill.txt", "r")
#take the name from the text file written by C
name = BillInfo.readline()
name = name.replace('\n', '') # delete any newline character in the string
name = name.replace('\r', '') # delete any enter character in the string

# take the pan from the text file written by the C
pan = BillInfo.readline()   
pan = pan.replace('\n', '') # delete any newline and enter characters in the string
pan = pan.replace('\r', '')

# take the date from the text file written by c
date = BillInfo.readline()   
date = date.replace('\n', '') # delete any newline and enter characters in the string
date = date.replace('\r', '')

# take the state of the transaction from the text file
tState = BillInfo.readline()
tState = tState.replace('\n', '')   # delete any newline or entre characters in the string
tState = tState.replace('\r', '')

# take the amount of money required in the transaction
amount_taken = BillInfo.readline()
amount_taken = amount_taken.replace('\n', '')   # delete any newline or enter in the string
amount_taken = amount_taken.replace('\r', '')

# take the remaining balance from the text file
remainingBalance = BillInfo.readline()
remainingBalance = remainingBalance.replace('\n', '')   # delete any newline or enter characters
remainingBalance = remainingBalance.replace('\r', '')

# take the sequence number from the text file
billSequenceNumber = BillInfo.readline() 
billSequenceNumber = billSequenceNumber.replace('\n', '') # delete any newline or enter characters
billSequenceNumber = billSequenceNumber.replace('\r', '')


#open the MS Word file
doc = Document("Database/Bill/Bill.docx")

# put the Name of the user
doc.paragraphs[1].runs[2].text = name

# put the PAN of the user
doc.paragraphs[2].runs[2].text = pan

# put the PAN of the user
doc.paragraphs[3].runs[2].text = date

# put Transaction status
doc.paragraphs[5].runs[0].text = tState
doc.paragraphs[5].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # align the text in the line to center

# put the Money taken of the user
doc.paragraphs[7].runs[2].text = amount_taken

# put the Balance remaining in User Account of the user
doc.paragraphs[8].runs[2].text = remainingBalance

# put the Sequence Number of the user
doc.paragraphs[10].runs[2].text = billSequenceNumber

# save the changes in the document then close the file and end the program
doc.save("Database/Bill/Bill.docx")
